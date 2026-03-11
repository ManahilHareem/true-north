"""
═══════════════════════════════════════════════════════════════════════════════
True North — Text Extraction Engine

Extracts raw text from uploaded files. Supports:
  - Plain text (.txt, .md, .csv, .log, .json, .xml, .yaml, .yml)
  - Code files (.py, .js, .ts, .jsx, .tsx)
  - PDF (.pdf) via PyPDF2
  - Word documents (.docx) via python-docx
  - HTML (.html, .htm) via basic tag stripping

Each extractor returns clean text ready for chunking and embedding.
═══════════════════════════════════════════════════════════════════════════════
"""

import re
from pathlib import Path
from typing import Optional


def extract_text(file_path: str, file_extension: str) -> str:
    """
    Extract text content from a file based on its extension.

    Args:
        file_path: Absolute path to the file.
        file_extension: File extension (e.g., ".pdf", ".txt").

    Returns:
        Extracted text as a string.

    Raises:
        ValueError: If the file type is unsupported.
    """
    ext = file_extension.lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    elif ext in {".html", ".htm"}:
        return _extract_html(file_path)
    elif ext in {
        ".txt", ".md", ".csv", ".json", ".log",
        ".xml", ".yaml", ".yml",
        ".py", ".js", ".ts", ".jsx", ".tsx",
    }:
        return _extract_plaintext(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(file_path: str) -> str:
    """Extract text from PDF using PyPDF2."""
    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append(f"--- Page {i + 1} ---\n{text.strip()}")
    return "\n\n".join(pages)


def _extract_docx(file_path: str) -> str:
    """Extract text from Word documents using python-docx."""
    from docx import Document

    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


def _extract_html(file_path: str) -> str:
    """Extract text from HTML by stripping tags."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    # Remove script and style blocks
    content = re.sub(r"<script[^>]*>.*?</script>", "", content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r"<style[^>]*>.*?</style>", "", content, flags=re.DOTALL | re.IGNORECASE)

    # Replace block-level tags with newlines
    content = re.sub(r"<(?:br|p|div|h[1-6]|li|tr)[^>]*>", "\n", content, flags=re.IGNORECASE)

    # Strip remaining tags
    content = re.sub(r"<[^>]+>", "", content)

    # Clean up whitespace
    content = re.sub(r"\n\s*\n+", "\n\n", content)
    return content.strip()


def _extract_plaintext(file_path: str) -> str:
    """Extract text from plain text files."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()
